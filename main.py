import streamlit as st
from datetime import datetime
import time
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2

# Page configuration
st.set_page_config(
    page_title="CaseLens - Home",
    page_icon="C",
    layout="wide"
)

# Custom CSS for CaseLens blue buttons
st.markdown("""
    <style>
    /* All buttons - filled with CaseLens blue */
    div[data-testid="stButton"] > button {
        background-color: #4D68F9 !important;
        color: white !important;
        border: none !important;
    }
    div[data-testid="stButton"] > button:hover {
        background-color: #3D58E9 !important;
        color: white !important;
    }
    </style>
    """, unsafe_allow_html=True)

# Initialize session state for navigation
if 'current_page' not in st.session_state:
    st.session_state.current_page = 'home'
if 'selected_case' not in st.session_state:
    st.session_state.selected_case = None
if 'processed_documents' not in st.session_state:
    st.session_state.processed_documents = []

# Sample case data (replace with your actual data)
cases = [
    {
        "id": 1,
        "name": "Hanessianadr Case 1",
        "description": "Harris FRC Acquisition vs RESEARCH CORPORATION TECHNOLOGIES",
        "documents": 156,
        "num_events": 3,
        "date_range": "1999-01-01 to 2025-09-30",
        "status": "Active",
        "last_updated": "2024-11-15"
    },
    {
        "id": 2,
        "name": "Patent Infringement Case 2",
        "description": "Technology patent dispute involving multiple parties",
        "documents": 243,
        "num_events": 5,
        "date_range": "2020-03-15 to 2025-06-30",
        "status": "Active",
        "last_updated": "2024-11-18"
    },
    {
        "id": 3,
        "name": "Contract Dispute Case 3",
        "description": "Commercial contract breach and damages claim",
        "documents": 89,
        "num_events": 2,
        "date_range": "2021-07-01 to 2024-12-31",
        "status": "Pending",
        "last_updated": "2024-11-10"
    },
    {
        "id": 4,
        "name": "Trademark Litigation Case 4",
        "description": "Brand trademark infringement proceedings",
        "documents": 312,
        "num_events": 8,
        "date_range": "2019-05-20 to 2025-08-15",
        "status": "Active",
        "last_updated": "2024-11-19"
    },
    {
        "id": 5,
        "name": "Arbitration Case 5",
        "description": "International arbitration dispute resolution",
        "documents": 178,
        "num_events": 4,
        "date_range": "2022-01-10 to 2025-11-30",
        "status": "In Review",
        "last_updated": "2024-11-12"
    }
]

def navigate_to_events(case):
    """Navigate to events page with selected case"""
    st.session_state.selected_case = case
    st.session_state.current_page = 'events'
    st.rerun()

def extract_text_from_pdf(pdf_file):
    """Extract text content from uploaded PDF"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return None

def simulate_ai_processing(pdf_text, pdf_name):
    """Simulate AI processing of the document (replace with actual AI logic)"""
    time.sleep(3)  # Simulate processing time
    
    # This is where you would call your actual AI service
    # For now, we'll create sample outputs
    
    summary = f"""
    Document Analysis Summary for: {pdf_name}
    
    This document has been analyzed and the following key points were identified:
    
    1. Main Topics:
       - Legal proceedings and case details
       - Important dates and timelines
       - Key parties involved
    
    2. Document Structure:
       - Total pages analyzed
       - Sections identified
       - Cross-references noted
    
    3. Key Findings:
       - Critical information extracted
       - Relevant citations identified
       - Important clauses highlighted
    
    Processing completed at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
    """
    
    detailed_analysis = f"""
    Detailed Analysis Report for: {pdf_name}
    
    EXECUTIVE SUMMARY:
    This report provides a comprehensive analysis of the uploaded document, 
    including extracted events, timeline construction, and relevant legal citations.
    
    DOCUMENT OVERVIEW:
    - File Name: {pdf_name}
    - Processing Date: {datetime.now().strftime('%Y-%m-%d')}
    - Analysis Type: Automated AI Processing
    
    EXTRACTED CONTENT:
    
    Section 1: Timeline of Events
    The document contains references to multiple events that have been chronologically 
    organized for case management purposes.
    
    Section 2: Key Entities and Parties
    Various parties and entities mentioned in the document have been identified and 
    categorized for reference.
    
    Section 3: Legal References
    Citations, statutes, and legal precedents mentioned in the document have been 
    extracted and organized.
    
    Section 4: Recommendations
    Based on the analysis, the following actions are recommended:
    - Review identified events for case timeline
    - Verify party information accuracy
    - Cross-reference with existing case documents
    
    CONCLUSION:
    The automated analysis has been completed successfully. Manual review is 
    recommended to verify AI-extracted information.
    
    Generated by: CaseLens AI Processing System
    Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
    """
    
    return summary, detailed_analysis

def create_word_document(title, content):
    """Create a Word document with the given content"""
    doc = Document()
    
    # Add title
    title_paragraph = doc.add_heading(title, level=1)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content
    for line in content.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(11)
    
    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def process_pdf_and_generate_docs(uploaded_file):
    """Process PDF and generate two Word documents"""
    # Extract text from PDF
    pdf_text = extract_text_from_pdf(uploaded_file)
    
    if pdf_text is None:
        return None, None
    
    # Simulate AI processing
    summary, detailed_analysis = simulate_ai_processing(pdf_text, uploaded_file.name)
    
    # Create Word documents
    doc1 = create_word_document("Summary Report", summary)
    doc2 = create_word_document("Detailed Analysis", detailed_analysis)
    
    return doc1, doc2

def show_home_page():
    """Display the home page with case selection"""
    
    # Sidebar with user info
    with st.sidebar:
        st.markdown("**User:** shushan@caselens.tech")
        st.divider()
        if st.button("📤 Upload Documents", use_container_width=True):
            st.session_state.current_page = 'upload'
            st.rerun()
        if st.button("Settings", use_container_width=True):
            st.session_state.current_page = 'settings'
            st.rerun()
    
    # Header
    st.markdown("### My Cases")
    st.divider()
    
    # Display cases in a grid layout (3 cards per row)
    for i in range(0, len(cases), 3):
        cols = st.columns(3)
        
        for j, col in enumerate(cols):
            if i + j < len(cases):
                case = cases[i + j]
                
                with col:
                    # Create a card using Streamlit native container with border
                    with st.container(border=True):
                        # Case name with View button
                        col_name, col_btn = st.columns([3, 1])
                        with col_name:
                            st.markdown(f"**{case['name']}**")
                        with col_btn:
                            if st.button("View", key=f"case_{case['id']}", type="secondary"):
                                navigate_to_events(case)
                        
                        # Case description with fixed height to keep cards uniform
                        st.markdown(f"""
                        <div style="height: 48px; overflow: hidden; text-overflow: ellipsis;">
                        <span style="color: #6c757d; font-size: 0.875rem;">{case['description']}</span>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        st.markdown("")  # Spacing
                        
                        # Information as colorful Streamlit native badges with labels
                        date_range_short = case['date_range'][:4] + '-' + case['date_range'][-4:]
                        
                        # Display tags with labels inline
                        col1, col2 = st.columns(2)
                        with col1:
                            st.markdown(f"**Status:** :blue-background[{case['status']}]")
                            st.markdown(f"**Documents:** :green-background[{case['documents']}]")
                        with col2:
                            st.markdown(f"**Events:** :orange-background[{case['num_events']}]")
                            st.markdown(f"**Period:** :gray-background[{date_range_short}]")

def show_events_page():
    """Display the events page for selected case"""
    if st.session_state.selected_case is None:
        st.error("No case selected. Returning to home page...")
        st.session_state.current_page = 'home'
        st.rerun()
        return
    
    case = st.session_state.selected_case
    
    # Sidebar with user info only
    with st.sidebar:
        st.markdown("**User:** shushan@caselens.tech")
    
    # Back button and header
    if st.button("← Back to Cases", type="secondary"):
        st.session_state.current_page = 'home'
        st.rerun()
    
    st.markdown(f"### {case['name']}")
    st.divider()
    
    # Tabs for different views
    tab1, tab2 = st.tabs(["Card View", "Table View"])
    
    with tab1:
        st.markdown("### Card View")
        st.info("This is where your events will be displayed in card format (similar to your second screenshot)")
        
        # Sample events display
        st.markdown("""
        **Event 1999-00-00**  
        In 1999, the definition of "Allowed Deductions" under the License Agreements between Harris FRC Acquisition, L.P. and RESEARCH CORPORATION TECHNOLOGIES, INC...
        
        1 Source
        """)
        
        st.divider()
        
        st.markdown("""
        **Event 2008-00-00**  
        From 2008 to 30 September 2015, a dispute in the ICC International Court of Arbitration (case number 27850/PDP) involves Harris FRC Acquisition...
        
        1 Source
        """)
    
    with tab2:
        st.markdown("### Table View")
        st.info("This is where your events will be displayed in table format")
        
        # Sample table
        import pandas as pd
        sample_data = pd.DataFrame({
            'Date': ['1999-00-00', '2008-00-00', '2008-01-01'],
            'Description': [
                'Definition of Allowed Deductions established',
                'ICC Arbitration case initiated',
                'Net sales reporting period begins'
            ],
            'Sources': [1, 1, 2]
        })
        st.dataframe(sample_data, use_container_width=True)

def show_upload_page():
    """Display the document upload page"""
    
    # Sidebar with user info
    with st.sidebar:
        st.markdown("**User:** shushan@caselens.tech")
        st.divider()
        if st.button("← Back to Cases", use_container_width=True):
            st.session_state.current_page = 'home'
            st.rerun()
    
    # Header
    st.markdown("### Upload Documents")
    st.divider()
    
    st.markdown("""
    Upload a PDF document to process. The AI will analyze the document and generate:
    - **Summary Report** - A concise overview of the document
    - **Detailed Analysis** - A comprehensive analysis with extracted information
    """)
    
    st.markdown("")
    
    # File uploader
    uploaded_file = st.file_uploader("Choose a PDF file", type=['pdf'], key="pdf_uploader")
    
    if uploaded_file is not None:
        st.success(f"✅ File uploaded: {uploaded_file.name}")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("🚀 Process Document", use_container_width=True, type="primary"):
                # Show progress
                with st.spinner('🤖 AI is processing your document...'):
                    progress_bar = st.progress(0)
                    
                    # Simulate processing steps
                    progress_bar.progress(20)
                    time.sleep(0.5)
                    st.info("📄 Reading PDF content...")
                    
                    progress_bar.progress(40)
                    time.sleep(1)
                    st.info("🧠 Analyzing document with AI...")
                    
                    progress_bar.progress(60)
                    
                    # Process the document
                    doc1, doc2 = process_pdf_and_generate_docs(uploaded_file)
                    
                    progress_bar.progress(80)
                    time.sleep(0.5)
                    st.info("📝 Generating Word documents...")
                    
                    progress_bar.progress(100)
                    time.sleep(0.5)
                
                if doc1 and doc2:
                    st.success("✅ Processing complete! Your documents are ready.")
                    
                    st.markdown("---")
                    st.markdown("### 📥 Download Your Documents")
                    
                    col_a, col_b = st.columns(2)
                    
                    with col_a:
                        st.markdown("**Summary Report**")
                        st.download_button(
                            label="📄 Download Summary",
                            data=doc1.getvalue(),
                            file_name=f"Summary_{uploaded_file.name.replace('.pdf', '')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    
                    with col_b:
                        st.markdown("**Detailed Analysis**")
                        st.download_button(
                            label="📄 Download Analysis",
                            data=doc2.getvalue(),
                            file_name=f"Analysis_{uploaded_file.name.replace('.pdf', '')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    
                    # Store in session state
                    st.session_state.processed_documents.append({
                        'filename': uploaded_file.name,
                        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    })
                else:
                    st.error("❌ Error processing document. Please try again.")
    
    # Show processing history
    if st.session_state.processed_documents:
        st.markdown("---")
        st.markdown("### 📋 Processing History")
        
        for idx, doc in enumerate(reversed(st.session_state.processed_documents)):
            with st.container(border=True):
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.markdown(f"**{doc['filename']}**")
                    st.caption(f"Processed: {doc['timestamp']}")
                with col2:
                    st.markdown("✅ Complete")

def show_settings_page():
    """Display the settings page"""
    
    col1, col2 = st.columns([6, 1])
    with col1:
        if st.button("← Back"):
            st.session_state.current_page = 'home'
            st.rerun()
        st.title("Settings")
    
    st.divider()
    
    st.markdown("### Account Settings")
    
    st.markdown("**Username / Email**")
    email = st.text_input("", value="shushan@caselens.tech", disabled=True)
    
    st.divider()
    
    if st.button("Log out", type="primary"):
        st.success("Logged out successfully!")

# Main app logic
def main():
    if st.session_state.current_page == 'home':
        show_home_page()
    elif st.session_state.current_page == 'events':
        show_events_page()
    elif st.session_state.current_page == 'upload':
        show_upload_page()
    elif st.session_state.current_page == 'settings':
        show_settings_page()

if __name__ == "__main__":
    main()
